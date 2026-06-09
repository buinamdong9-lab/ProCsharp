from pathlib import Path
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "BaoCao"
ASSETS = OUT / "assets"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)

DOCX_PATH = OUT / "Bao_cao_ky_thuat_He_thong_quan_ly_muon_tra_thiet_bi_CO_PHAN_TICH_DU_LIEU.docx"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF4FA"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
GREEN = "2E7D32"
AMBER = "A66A00"
RED = "A61B1B"
BLACK = "000000"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B7C9DB", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_widths(table, widths):
    table.autofit = False
    total_dxa = sum(int(w * 1440) for w in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            dxa = int(width * 1440)
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def font(run, name="Times New Roman", size=13, bold=None, italic=None, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.first_line_indent = Inches(0.3)

    for name, size, color, before, after in (
        ("Title", 25, NAVY, 0, 8),
        ("Subtitle", 15, GRAY, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 14, BLUE, 12, 6),
        ("Heading 3", 13, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Times New Roman"
        style.font.size = Pt(13)
        style.paragraph_format.left_indent = Inches(0.45)
        style.paragraph_format.first_line_indent = Inches(-0.22)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    font(run, size=10, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(0.85)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run("BÁO CÁO DỰ ÁN KỸ THUẬT LẬP TRÌNH")
        font(hr, size=9, bold=True, color=GRAY)
        add_page_number(section.footer.paragraphs[0])


def add_para(doc, text="", bold_lead=None, align=None, first_indent=True, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if not first_indent:
        p.paragraph_format.first_line_indent = Inches(0)
    if keep:
        p.paragraph_format.keep_with_next = True
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        font(r2)
    else:
        r = p.add_run(text)
        font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = Inches(-0.22)
        r = p.add_run(item)
        font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.first_line_indent = Inches(-0.22)
        r = p.add_run(item)
        font(r)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Inches(0)
    return p


def add_table(doc, headers, rows, widths, font_size=10.5, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        font(r, size=font_size, bold=True, color=NAVY)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for c in cells:
                set_cell_shading(c, "F9FBFD")
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            font(r, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, body, fill=PALE_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    set_table_widths(table, [6.15])
    set_table_borders(table, color="B7C9DB", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    font(r, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    font(r2, size=12)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    set_table_widths(table, [6.15])
    set_table_borders(table, color="D0D5DD", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code.splitlines()):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        font(r, name="Consolas", size=9.2, color="333333")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Mục lục sẽ được cập nhật khi mở tài liệu."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, placeholder, fld_end])


def diagram(path, title, nodes, arrows):
    width, height = 1500, 850
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_path = r"C:\Windows\Fonts\arial.ttf"
    bold_path = r"C:\Windows\Fonts\arialbd.ttf"
    f_title = ImageFont.truetype(bold_path, 44)
    f_node = ImageFont.truetype(bold_path, 25)
    f_small = ImageFont.truetype(font_path, 20)
    draw.text((width // 2, 40), title, fill=(23, 54, 93), anchor="ma", font=f_title)
    positions = {}
    for key, x, y, w, h, label, fill in nodes:
        positions[key] = (x, y, w, h)
        draw.rounded_rectangle((x, y, x + w, y + h), radius=20, fill=fill, outline=(46, 116, 181), width=4)
        lines = label.split("\n")
        total_h = len(lines) * 31
        start_y = y + (h - total_h) / 2
        for idx, line in enumerate(lines):
            draw.text((x + w / 2, start_y + idx * 31), line, fill=(23, 54, 93), anchor="ma", font=f_node if idx == 0 else f_small)
    for source, target, label in arrows:
        sx, sy, sw, sh = positions[source]
        tx, ty, tw, th = positions[target]
        start = (sx + sw / 2, sy + sh / 2)
        end = (tx + tw / 2, ty + th / 2)
        dx, dy = end[0] - start[0], end[1] - start[1]
        if abs(dx) > abs(dy):
            start = (sx + (sw if dx > 0 else 0), sy + sh / 2)
            end = (tx + (0 if dx > 0 else tw), ty + th / 2)
        else:
            start = (sx + sw / 2, sy + (sh if dy > 0 else 0))
            end = (tx + tw / 2, ty + (0 if dy > 0 else th))
        draw.line((start, end), fill=(70, 85, 105), width=5)
        import math
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        length = 18
        a1 = angle + 2.55
        a2 = angle - 2.55
        p1 = (end[0] + length * math.cos(a1), end[1] + length * math.sin(a1))
        p2 = (end[0] + length * math.cos(a2), end[1] + length * math.sin(a2))
        draw.polygon([end, p1, p2], fill=(70, 85, 105))
        if label:
            mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
            bbox = draw.textbbox((mx, my), label, anchor="mm", font=f_small)
            draw.rectangle((bbox[0] - 6, bbox[1] - 3, bbox[2] + 6, bbox[3] + 3), fill="white")
            draw.text((mx, my), label, fill=(70, 85, 105), anchor="mm", font=f_small)
    img.save(path)


def create_diagrams():
    diagram(
        ASSETS / "architecture.png",
        "Kiến trúc tổng thể của hệ thống",
        [
            ("ui", 550, 125, 400, 110, "GUI - Windows Forms\nForm và UserControl", (225, 238, 249)),
            ("bll", 550, 335, 400, 120, "BLL - Dịch vụ nghiệp vụ\nPhê duyệt mượn / trả", (226, 240, 217)),
            ("dal", 550, 565, 400, 120, "DAL - Repository\nADO.NET, truy vấn tham số", (255, 239, 204)),
            ("db", 1100, 565, 300, 120, "SQL Server\nQuanLyThietBi", (245, 222, 222)),
            ("log", 100, 335, 300, 120, "Hạ tầng dùng chung\nValidation, logging", (237, 230, 247)),
        ],
        [
            ("ui", "bll", "gọi nghiệp vụ"),
            ("bll", "dal", "đọc/ghi dữ liệu"),
            ("dal", "db", "Microsoft.Data.SqlClient"),
            ("ui", "log", "kiểm tra / ghi log"),
            ("log", "bll", "hỗ trợ"),
        ],
    )
    diagram(
        ASSETS / "state.png",
        "Vòng đời phiếu mượn",
        [
            ("pending", 120, 170, 280, 100, "PENDING\nChờ duyệt", (255, 244, 204)),
            ("borrowing", 590, 170, 300, 100, "BORROWING\nĐang mượn", (220, 235, 250)),
            ("return_pending", 1080, 170, 300, 100, "RETURN_PENDING\nChờ duyệt trả", (230, 222, 247)),
            ("returned", 1080, 520, 300, 100, "RETURNED\nĐã trả", (222, 239, 217)),
            ("rejected", 120, 520, 280, 100, "REJECTED\nTừ chối", (245, 222, 222)),
        ],
        [
            ("pending", "borrowing", "duyệt mượn"),
            ("pending", "rejected", "từ chối"),
            ("borrowing", "return_pending", "gửi yêu cầu trả"),
            ("return_pending", "returned", "duyệt trả đủ"),
            ("return_pending", "borrowing", "từ chối / trả một phần"),
        ],
    )
    diagram(
        ASSETS / "erd.png",
        "Mô hình dữ liệu khái quát",
        [
            ("roles", 90, 120, 260, 90, "Roles", (225, 238, 249)),
            ("users", 90, 350, 260, 100, "Users", (225, 238, 249)),
            ("tickets", 520, 350, 310, 100, "BorrowTickets", (226, 240, 217)),
            ("details", 1040, 350, 310, 100, "BorrowDetails", (255, 239, 204)),
            ("devices", 1040, 600, 310, 100, "Devices", (245, 222, 222)),
            ("instances", 1040, 105, 310, 100, "DeviceInstances", (245, 222, 222)),
            ("rooms", 520, 600, 310, 100, "Rooms", (237, 230, 247)),
            ("returns", 520, 105, 310, 100, "ReturnRequests\n+ Details", (230, 222, 247)),
        ],
        [
            ("roles", "users", "1 - n"),
            ("users", "tickets", "1 - n"),
            ("tickets", "details", "1 - n"),
            ("devices", "details", "1 - n"),
            ("instances", "details", "0..1 - n"),
            ("rooms", "devices", "1 - n"),
            ("tickets", "returns", "1 - 0..1"),
        ],
    )
    diagram(
        ASSETS / "borrow_flow.png",
        "Luồng lập và phê duyệt phiếu mượn",
        [
            ("select", 70, 310, 250, 100, "Chọn người mượn\nvà cá thể", (225, 238, 249)),
            ("validate", 390, 310, 250, 100, "Kiểm tra hạn trả,\ngiới hạn, trùng cá thể", (255, 239, 204)),
            ("save", 710, 310, 250, 100, "Transaction tạo\nphiếu PENDING", (226, 240, 217)),
            ("approve", 1030, 310, 250, 100, "Admin duyệt\nvà khóa tồn kho", (230, 222, 247)),
            ("update", 1320, 310, 150, 100, "BORROWING", (222, 239, 217)),
        ],
        [
            ("select", "validate", ""),
            ("validate", "save", ""),
            ("save", "approve", ""),
            ("approve", "update", ""),
        ],
    )


def add_figure(doc, path, caption, width=6.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", caption)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.paragraph_format.first_line_indent = Inches(0)


def cover(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("VIỆN CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG")
    font(r, size=14, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("BỘ MÔN CÔNG NGHỆ PHẦN MỀM")
    font(r, size=14, bold=True, color=NAVY)
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("BÁO CÁO BÀI TẬP LỚN")
    font(r, size=25, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("HỌC PHẦN: DỰ ÁN KỸ THUẬT LẬP TRÌNH")
    font(r, size=16, bold=True, color=BLUE)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("ĐỀ TÀI")
    font(r, size=14, bold=True, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("HỆ THỐNG QUẢN LÝ MƯỢN TRẢ THIẾT BỊ")
    font(r, size=23, bold=True, color=NAVY)
    for _ in range(3):
        doc.add_paragraph()
    add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ["Họ tên học viên", "Bùi Nam Đông"],
            ["Đơn vị", "Lớp TTNT, Đại đội 160"],
            ["Giảng viên", "3// Tống Minh Đức"],
            ["Công nghệ", "C# WinForms, .NET 8, SQL Server Express"],
        ],
        [1.8, 4.35],
        font_size=12,
        header_fill="E8EEF5",
    )
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("Hà Nội, tháng 6 năm 2026")
    font(r, size=13, bold=True, color=NAVY)
    doc.add_page_break()


def front_matter(doc):
    add_heading(doc, "LỜI CAM ĐOAN", 1)
    add_para(doc, "Tôi cam đoan báo cáo này phản ánh đúng quá trình khảo sát, phân tích và rà soát mã nguồn của dự án Hệ thống quản lý mượn trả thiết bị. Các nhận định kỹ thuật được đối chiếu với mã nguồn C#, cấu hình dự án và cơ sở dữ liệu SQL Server tại thời điểm ngày 05/06/2026. Kết quả build, số lượng cảnh báo và số liệu kiểm tra dữ liệu chỉ được ghi nhận khi có lệnh xác minh tương ứng.")
    add_para(doc, "Báo cáo phân biệt rõ chức năng đã hiện thực, chức năng mới dừng ở giao diện/cấu hình và nội dung đề xuất phát triển. Học viên chịu trách nhiệm về tính trung thực của nội dung.")
    add_para(doc, "Bùi Nam Đông", align=WD_ALIGN_PARAGRAPH.RIGHT, first_indent=False)
    doc.add_page_break()

    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(doc, "Trong quá trình thực hiện dự án, tôi đã nhận được sự hướng dẫn về phương pháp phân tích bài toán, tổ chức chương trình, thiết kế cơ sở dữ liệu và kiểm thử sản phẩm. Tôi xin trân trọng cảm ơn giảng viên 3// Tống Minh Đức đã định hướng, góp ý và tạo điều kiện để tôi hoàn thành bài tập lớn.")
    add_para(doc, "Tôi cũng cảm ơn các đồng chí trong lớp đã trao đổi về yêu cầu nghiệp vụ và hỗ trợ kiểm tra giao diện. Do phạm vi dự án khá rộng, báo cáo khó tránh khỏi thiếu sót; tôi kính mong tiếp tục nhận được ý kiến góp ý để sản phẩm được hoàn thiện hơn.")
    doc.add_page_break()

    add_heading(doc, "TÓM TẮT BÁO CÁO", 1)
    add_para(doc, "Dự án xây dựng ứng dụng desktop quản lý mượn trả thiết bị cho môi trường lớp học hoặc đơn vị đào tạo. Hệ thống được phát triển bằng C# Windows Forms trên .NET 8, sử dụng SQL Server Express và thư viện Microsoft.Data.SqlClient. Phạm vi chức năng gồm đăng nhập, khóa tài khoản khi đăng nhập sai, phân quyền theo vai trò, quản lý thiết bị và cá thể, quản lý phòng học, quản lý người dùng, lập phiếu mượn, phê duyệt mượn, gửi yêu cầu trả, phê duyệt trả toàn phần hoặc một phần, thùng rác, dashboard, báo cáo và xuất Excel.")
    add_para(doc, "Mã nguồn có 60 tệp C# với khoảng 15.240 dòng, tổ chức theo các thư mục GUI, BLL, DAL, Models và DataStructures. Dữ liệu được chuẩn hóa thành các bảng Users, Roles, Rooms, Devices, DeviceInstances, BorrowTickets, BorrowDetails, ReturnRequests và ReturnRequestDetails. Những thao tác nhạy cảm như duyệt mượn, duyệt trả, tạo phiếu và xóa vĩnh viễn được bao bọc bằng transaction.")
    add_para(doc, "Kết quả xác minh ngày 05/06/2026: build Release thành công, 0 lỗi và 56 cảnh báo; chưa có dự án kiểm thử tự động. Cơ sở dữ liệu thực tế có 108 người dùng, 38 thiết bị, 1.050 cá thể, 24 phòng, 40 phiếu và 66 dòng chi tiết. Kiểm tra chỉ đọc không phát hiện số lượng âm, khóa ngoại mồ côi hoặc mã thiết bị/cá thể trùng. Tuy nhiên, 30 dòng mượn đang hoạt động là dữ liệu lịch sử chưa gắn InstanceID, dẫn đến số lượng tổng hợp và trạng thái từng cá thể chưa hoàn toàn đồng bộ.")
    add_callout(doc, "Kết luận ngắn", "Sản phẩm có phạm vi nghiệp vụ tốt và đã xử lý nhiều tình huống thực tế hơn một ứng dụng CRUD cơ bản. Điểm cần ưu tiên là chuẩn hóa migration, bổ sung test tự động, xử lý dữ liệu lịch sử không có InstanceID, giảm cảnh báo nullable và tách trách nhiệm trong các lớp giao diện lớn.")

    add_heading(doc, "DANH MỤC TỪ VIẾT TẮT", 1)
    add_table(doc, ["Từ viết tắt", "Tiếng Anh", "Ý nghĩa"], [
        ["GUI", "Graphical User Interface", "Tầng giao diện Windows Forms."],
        ["BLL", "Business Logic Layer", "Tầng xử lý quy tắc mượn, trả và phê duyệt."],
        ["DAL", "Data Access Layer", "Tầng truy cập SQL Server qua repository."],
        ["CRUD", "Create - Read - Update - Delete", "Thêm, đọc, sửa, xóa dữ liệu."],
        ["RBAC", "Role-Based Access Control", "Phân quyền dựa trên vai trò."],
        ["DTO", "Data Transfer Object", "Đối tượng chuyển dữ liệu giữa các lớp."],
        ["ADO.NET", "ActiveX Data Objects for .NET", "Nền tảng truy cập dữ liệu của .NET."],
        ["SQL", "Structured Query Language", "Ngôn ngữ thao tác cơ sở dữ liệu."],
        ["FK", "Foreign Key", "Khóa ngoại bảo đảm quan hệ dữ liệu."],
        ["UI", "User Interface", "Giao diện người dùng."],
    ], [1.0, 2.2, 2.95])
    doc.add_page_break()

    add_heading(doc, "MỤC LỤC", 1)
    add_toc(doc)
    doc.add_page_break()


def chapter_intro(doc):
    add_heading(doc, "PHẦN I. MỞ ĐẦU", 1)
    add_heading(doc, "1. Lý do chọn đề tài", 2)
    add_para(doc, "Thiết bị phục vụ giảng dạy và làm việc thường được phân bố ở nhiều phòng, có số lượng lớn và có vòng đời sử dụng khác nhau. Nếu chỉ theo dõi bằng sổ tay hoặc bảng tính rời rạc, đơn vị khó xác định thiết bị nào đang sẵn sàng, cá thể nào đang được mượn, ai chịu trách nhiệm, phiếu nào quá hạn và tình trạng thiết bị khi trả.")
    add_para(doc, "Đề tài Hệ thống quản lý mượn trả thiết bị được lựa chọn nhằm giải quyết đồng thời ba nhóm vấn đề: quản lý danh mục tài sản, kiểm soát quy trình mượn trả và cung cấp thông tin quản trị. Đây là bài toán phù hợp với học phần Kỹ thuật lập trình vì yêu cầu kết hợp lập trình hướng đối tượng, giao diện sự kiện, xử lý dữ liệu, transaction, phân quyền, validation và báo cáo.")

    add_heading(doc, "2. Mục tiêu", 2)
    add_bullets(doc, [
        "Xây dựng ứng dụng desktop thống nhất cho quản lý thiết bị, cá thể, phòng học và người dùng.",
        "Kiểm soát vòng đời phiếu từ chờ duyệt, đang mượn, chờ duyệt trả đến đã trả hoặc bị từ chối.",
        "Bảo đảm không xuất kho khi thiết bị hết số lượng, đang bảo trì, ngừng sử dụng hoặc cá thể không ở tình trạng tốt.",
        "Phân quyền rõ giữa người dùng thông thường, thủ kho/nhân viên và quản trị viên.",
        "Theo dõi quá hạn, thiết bị có vấn đề, thống kê theo thời gian và xuất báo cáo Excel.",
        "Tăng tính an toàn bằng password hashing, tạm khóa đăng nhập, transaction, xóa mềm, thùng rác và logging.",
        "Rà soát kỹ thuật để xác định rủi ro hiện tại và đề xuất lộ trình phát triển có thể thực hiện.",
    ])

    add_heading(doc, "3. Đối tượng và phạm vi", 2)
    add_para(doc, "Đối tượng quản lý gồm loại thiết bị, bản ghi thiết bị tổng hợp, từng cá thể có mã tài sản, phòng học, người dùng, vai trò, phiếu mượn, chi tiết mượn, yêu cầu trả và chi tiết trả. Ứng dụng hướng đến môi trường Windows trong mạng nội bộ hoặc trên máy có kết nối tới SQL Server Express.")
    add_para(doc, "Phạm vi hiện tại chưa bao gồm gửi email nhắc hạn, lập lịch nền, tích hợp mã QR/barcode, chữ ký số, API web, đồng bộ đa chi nhánh hoặc sao lưu cơ sở dữ liệu tự động. Giao diện cấu hình có trường liên quan email nhưng mã nguồn chủ động khóa chức năng và ghi rõ chưa triển khai.")

    add_heading(doc, "4. Phương pháp thực hiện và nguồn bằng chứng", 2)
    add_numbered(doc, [
        "Đọc toàn bộ cấu trúc repository và thống kê quy mô mã nguồn.",
        "Đối chiếu luồng sự kiện từ Windows Forms đến BLL, DAL và SQL Server.",
        "Build cấu hình Release bằng lệnh dotnet build FrmProject.slnx -c Release.",
        "Kiểm tra schema, số lượng bản ghi, khóa ngoại, chỉ mục, trạng thái và tính nhất quán bằng truy vấn SQL chỉ đọc.",
        "Phân loại phát hiện review theo mức nghiêm trọng, tác động và khuyến nghị.",
        "Soạn báo cáo theo cấu trúc học thuật, tránh khẳng định các chức năng chưa có bằng chứng.",
    ])


def chapter_requirements(doc):
    add_heading(doc, "PHẦN II. NỘI DUNG", 1)
    add_heading(doc, "CHƯƠNG 1. KHẢO SÁT VÀ PHÂN TÍCH YÊU CẦU", 1)
    add_heading(doc, "1.1. Bài toán nghiệp vụ", 2)
    add_para(doc, "Đơn vị cần biết tổng số thiết bị, số đang cho mượn, số quá hạn, số đang bảo trì và lịch sử thao tác. Một thiết bị tổng hợp có thể có nhiều cá thể. Mỗi cá thể có mã tài sản riêng, trạng thái và tình trạng vật lý riêng. Phiếu mượn phải gắn với người mượn, thời gian, mục đích, phòng sử dụng và các cá thể cụ thể.")
    add_para(doc, "Quy trình không chỉ là trừ/cộng số lượng. Trước khi duyệt mượn, hệ thống phải khóa và kiểm tra lại dữ liệu để tránh hai người cùng lấy một cá thể. Khi trả, người dùng có thể trả đủ, trả một phần, báo hỏng, mất hoặc thiếu phụ kiện. Quản trị viên cần duyệt kết quả trả để cập nhật đúng tồn kho và tình trạng.")

    add_heading(doc, "1.2. Tác nhân và quyền", 2)
    add_table(doc, ["Tác nhân", "Mục tiêu", "Quyền chính"], [
        ["Người dùng", "Tạo và theo dõi giao dịch của chính mình.", "Lập phiếu mượn, gửi yêu cầu trả, xem danh sách phiếu cá nhân."],
        ["Thủ kho/Nhân viên", "Vận hành kho và quy trình mượn trả.", "Dashboard, thiết bị, phòng, phiếu mượn/trả, danh sách phiếu, báo cáo."],
        ["Quản trị viên", "Quản trị toàn hệ thống.", "Toàn bộ quyền của thủ kho, cộng quản lý người dùng, cấu hình và thùng rác."],
    ], [1.25, 2.15, 2.75])
    add_para(doc, "Vai trò được ánh xạ linh hoạt bằng DbSchemaHelper.ResolveRole. Chuỗi chứa “admin” hoặc “quản trị” được coi là Admin; “thủ kho”, “quản lý” hoặc “staff” được coi là Staff; các vai trò còn lại như Sinh viên, Giảng viên được coi là User.")

    add_heading(doc, "1.3. Yêu cầu chức năng", 2)
    fr_rows = [
        ["FR-01", "Đăng nhập", "Xác minh username và mật khẩu băm; tạm khóa sau 5 lần sai trong 5 phút."],
        ["FR-02", "Phân quyền", "Ẩn/hiện menu và giới hạn dữ liệu phiếu theo vai trò."],
        ["FR-03", "Thiết bị", "Thêm, sửa, lọc, phân trang, xóa mềm và khôi phục thiết bị."],
        ["FR-04", "Cá thể", "Sinh mã tài sản, theo dõi trạng thái/tình trạng, đồng bộ số lượng tổng."],
        ["FR-05", "Phòng học", "Quản lý mã phòng, tên, loại, tầng, sức chứa và trạng thái."],
        ["FR-06", "Người dùng", "Quản lý hồ sơ, vai trò, trạng thái, mật khẩu và xuất Excel."],
        ["FR-07", "Lập phiếu", "Chọn người mượn, phòng, cá thể, hạn trả và giới hạn số lượng."],
        ["FR-08", "Duyệt mượn", "Kiểm tra trạng thái/tồn kho trong transaction rồi chuyển sang BORROWING."],
        ["FR-09", "Yêu cầu trả", "Gửi danh sách trả đủ/thiếu/lỗi và chuyển RETURN_PENDING."],
        ["FR-10", "Duyệt trả", "Cộng tồn, cập nhật tình trạng, hỗ trợ duyệt toàn bộ hoặc một phần."],
        ["FR-11", "Danh sách phiếu", "Tìm kiếm, lọc trạng thái, xem chi tiết và thống kê nhanh."],
        ["FR-12", "Dashboard", "Hiển thị tổng thiết bị, đang mượn, quá hạn, sự cố và dữ liệu 6 tháng."],
        ["FR-13", "Báo cáo", "Thống kê tháng, top thiết bị, phiếu quá hạn và xuất XLSX."],
        ["FR-14", "Thùng rác", "Khôi phục hoặc xóa vĩnh viễn dữ liệu đã xóa mềm có kiểm tra tham chiếu."],
        ["FR-15", "Cấu hình", "Thiết lập hạn mượn, số thiết bị tối đa, yêu cầu phê duyệt và thông tin đơn vị."],
        ["FR-16", "Logging", "Ghi log ứng dụng, lỗi chưa xử lý và lịch sử đăng nhập."],
    ]
    add_table(doc, ["Mã", "Nhóm", "Mô tả"], fr_rows, [0.65, 1.35, 4.15], font_size=9.8)

    add_heading(doc, "1.4. Yêu cầu phi chức năng", 2)
    add_table(doc, ["Mã", "Thuộc tính", "Tiêu chí"], [
        ["NFR-01", "Đúng đắn", "Không cho số lượng âm; số trả không vượt số đang mượn; ngày trả dự kiến không trước ngày mượn."],
        ["NFR-02", "An toàn dữ liệu", "Transaction cho nghiệp vụ nhiều bước; khóa cập nhật khi duyệt mượn; FK và check constraint."],
        ["NFR-03", "Bảo mật", "Mật khẩu dùng ASP.NET Core Identity PasswordHasher; khóa đăng nhập; phân quyền theo vai trò."],
        ["NFR-04", "Hiệu năng", "Phân trang phía server cho thiết bị/người dùng/dashboard; truy vấn tham số; tải dashboard nền."],
        ["NFR-05", "Khả dụng", "Giao diện tiếng Việt, thông báo lỗi, xác nhận thao tác phá hủy, bộ lọc và điều hướng."],
        ["NFR-06", "Bảo trì", "Tách thư mục theo trách nhiệm; hằng số trạng thái; helper validation, schema và logging."],
        ["NFR-07", "Truy vết", "Log file theo ngày và thông tin trạng thái phiếu trong cơ sở dữ liệu."],
        ["NFR-08", "Tương thích", "Windows, .NET 8 Desktop Runtime, SQL Server Express, độ phân giải desktop."],
    ], [0.75, 1.35, 4.05], font_size=10)


def chapter_usecase(doc):
    add_heading(doc, "CHƯƠNG 2. USE CASE VÀ LUỒNG NGHIỆP VỤ", 1)
    add_heading(doc, "2.1. Danh sách use case", 2)
    add_table(doc, ["Mã", "Use case", "Tác nhân", "Kết quả"], [
        ["UC-01", "Đăng nhập", "Tất cả", "Mở giao diện đúng quyền hoặc ghi nhận thất bại."],
        ["UC-02", "Quản lý thiết bị/cá thể", "Staff, Admin", "Danh mục và số lượng được cập nhật."],
        ["UC-03", "Lập phiếu mượn", "User, Staff, Admin", "Tạo phiếu PENDING hoặc duyệt ngay theo cấu hình."],
        ["UC-04", "Duyệt/từ chối mượn", "Staff, Admin", "Chuyển BORROWING hoặc REJECTED."],
        ["UC-05", "Gửi yêu cầu trả", "User, Staff, Admin", "Lưu ReturnRequest và chuyển RETURN_PENDING."],
        ["UC-06", "Duyệt/từ chối trả", "Staff, Admin", "Cập nhật tồn kho và trạng thái phiếu."],
        ["UC-07", "Tra cứu phiếu", "Tất cả", "Danh sách đúng phạm vi quyền và trạng thái."],
        ["UC-08", "Xem báo cáo", "Staff, Admin", "Bảng thống kê và tệp Excel."],
        ["UC-09", "Khôi phục dữ liệu", "Admin", "Khôi phục bản ghi xóa mềm nếu thỏa điều kiện."],
    ], [0.7, 1.75, 1.2, 2.5], font_size=10)

    add_heading(doc, "2.2. Đặc tả UC-03: Lập phiếu mượn", 2)
    add_table(doc, ["Mục", "Nội dung"], [
        ["Tiền điều kiện", "Người dùng đã đăng nhập; cá thể ở trạng thái Có sẵn và tình trạng Tốt; phòng chưa ngừng sử dụng."],
        ["Luồng chính", "Chọn người mượn, phòng, thiết bị, cá thể và hạn trả; thêm cá thể vào danh sách; lưu phiếu."],
        ["Kiểm tra", "Không trùng cá thể; không vượt số thiết bị tối đa; hạn trả không trước ngày mượn."],
        ["Hậu điều kiện", "Tạo BorrowTickets và BorrowDetails trong transaction. Trạng thái ban đầu PENDING."],
        ["Ngoại lệ", "Schema thiếu cột, mã phiếu trùng, dữ liệu bị thay đổi đồng thời hoặc mất kết nối."],
    ], [1.35, 4.8], font_size=10.5)
    add_figure(doc, ASSETS / "borrow_flow.png", "Hình 2.1. Luồng lập và phê duyệt phiếu mượn")

    add_heading(doc, "2.3. Đặc tả UC-06: Duyệt trả", 2)
    add_para(doc, "Dịch vụ ReturnApprovalService tải yêu cầu trả có cấu trúc, khóa các dòng BorrowDetails cần kiểm tra, xác minh số lượng trả không vượt số đang mượn, cập nhật trạng thái cá thể và số lượng tổng hợp. Với thiết bị trả tốt, cá thể chuyển về Có sẵn; nếu hỏng hoặc có vấn đề, cá thể chuyển Bảo trì; nếu thiết bị cha đã ngừng sử dụng, cá thể giữ trạng thái Ngừng sử dụng.")
    add_para(doc, "Hệ thống hỗ trợ duyệt một phần. Các dòng đã duyệt được xóa khỏi ReturnRequestDetails, phần còn lại tiếp tục ở RETURN_PENDING. Khi toàn bộ số lượng đã trả, phiếu chuyển RETURNED và ghi ngày/người duyệt trả. Nếu yêu cầu bị từ chối, phiếu quay lại BORROWING.")

    add_heading(doc, "2.4. Máy trạng thái phiếu", 2)
    add_figure(doc, ASSETS / "state.png", "Hình 2.2. Các trạng thái và chuyển tiếp chính của phiếu")
    add_callout(doc, "Quy tắc bất biến", "Một phiếu chỉ được duyệt mượn khi đang PENDING; chỉ được gửi yêu cầu trả khi đang BORROWING; chỉ được duyệt hoặc từ chối trả khi đang RETURN_PENDING. Các lệnh UPDATE đều kiểm tra trạng thái hiện tại để tránh thao tác lặp.")


def chapter_architecture(doc):
    add_heading(doc, "CHƯƠNG 3. KIẾN TRÚC VÀ TỔ CHỨC MÃ NGUỒN", 1)
    add_heading(doc, "3.1. Kiến trúc tổng thể", 2)
    add_figure(doc, ASSETS / "architecture.png", "Hình 3.1. Kiến trúc logic của ứng dụng")
    add_para(doc, "Dự án áp dụng kiến trúc phân lớp ở mức thư mục. GUI chứa Form và UserControl; BLL chứa các dịch vụ phê duyệt và helper nghiệp vụ; DAL chứa repository và helper schema; Models chứa các đối tượng dữ liệu; DataStructures chứa validation, password hashing, logging và helper giao diện.")
    add_para(doc, "Kiến trúc chưa phải ba lớp thuần túy vì GUI gọi trực tiếp nhiều repository, còn BLL cũng truy cập SqlConnection. Tuy nhiên, cách chia hiện tại vẫn giảm đáng kể việc viết SQL trực tiếp trong event handler và tạo nền tảng để tiếp tục tách service/interface.")

    add_heading(doc, "3.2. Cấu trúc thư mục", 2)
    add_table(doc, ["Thư mục", "Vai trò", "Ví dụ"], [
        ["GUI", "Giao diện, sự kiện, binding DataGridView, điều hướng.", "FrmLogin, FrmMain, UcLapthietbi, UcTrathietbi."],
        ["BLL", "Quy tắc nghiệp vụ nhiều bước và transaction.", "BorrowApprovalService, ReturnApprovalService."],
        ["DAL", "Kết nối, truy vấn, repository, migration thích nghi schema.", "DbHelper, DbSchemaHelper, các Repository."],
        ["Models", "DTO, trạng thái, dữ liệu trao đổi.", "BorrowTicketDraft, AuthModels, AppStatuses."],
        ["DataStructures", "Tiện ích dùng chung.", "PasswordHelper, ValidationHelper, AppLogger."],
        ["Data", "Tài nguyên và notebook schema cũ.", "bangdatabase.ipynb, Resources."],
    ], [1.15, 2.75, 2.25], font_size=10.2)

    add_heading(doc, "3.3. Các module trọng yếu", 2)
    add_table(doc, ["Module", "Trách nhiệm", "Nhận xét review"], [
        ["FrmLogin/AuthRepository", "Xác thực, tạm khóa, trạng thái tài khoản.", "Có xử lý cả username không tồn tại để giảm brute-force trong phiên chạy."],
        ["FrmMain", "Điều hướng, phân quyền, dashboard, cache UserControl.", "584 dòng; đang gánh nhiều trách nhiệm UI và tải dữ liệu."],
        ["BorrowTicketRepository", "Tạo phiếu PENDING và chi tiết.", "Transaction tốt; schema detection linh hoạt nhưng tăng độ phức tạp."],
        ["BorrowApprovalService", "Khóa dữ liệu, kiểm tra tồn, trừ kho, đổi trạng thái cá thể.", "Sử dụng UPDLOCK/HOLDLOCK giúp giảm race condition."],
        ["ReturnApprovalService", "Duyệt trả đủ/một phần, cập nhật tình trạng.", "473 dòng; nghiệp vụ đúng hướng nhưng cần tách nhỏ để dễ test."],
        ["RecycleBinRepository", "Khôi phục/xóa vĩnh viễn.", "Có kiểm tra tham chiếu trước xóa, dùng transaction."],
        ["ReportRepository", "Báo cáo tháng, top thiết bị, quá hạn.", "Truy vấn rõ; FORMAT trong GROUP BY có thể chậm khi dữ liệu lớn."],
    ], [1.55, 2.45, 2.15], font_size=9.6)

    add_heading(doc, "3.4. Luồng khởi động", 2)
    add_numbered(doc, [
        "ApplicationConfiguration.Initialize cấu hình môi trường WinForms.",
        "AppLogger.Initialize tạo logger Serilog ghi file theo ngày.",
        "Đăng ký xử lý ThreadException và UnhandledException.",
        "Mở kết nối và chạy EnsureReturnedQuantitySchemaAndRestore để bổ sung schema cần thiết.",
        "Hiển thị FrmLogin; sau khi đăng nhập mở FrmMain theo vai trò.",
        "Khi kết thúc, CloseAndFlush bảo đảm log được ghi hết.",
    ])
    add_code(doc, """dotnet restore FrmProject.slnx
dotnet build FrmProject.slnx -c Release
dotnet run --project FrmProject\\FrmProject.csproj""")


def chapter_database(doc):
    add_heading(doc, "CHƯƠNG 4. THIẾT KẾ CƠ SỞ DỮ LIỆU", 1)
    add_heading(doc, "4.1. Mô hình quan hệ", 2)
    add_figure(doc, ASSETS / "erd.png", "Hình 4.1. Các thực thể và quan hệ chính")
    add_para(doc, "Dữ liệu được lưu trong database QuanLyThietBi trên SQL Server Express. Mô hình tách thiết bị tổng hợp khỏi từng cá thể: Devices lưu tên, loại, phòng và số lượng; DeviceInstances lưu mã tài sản, trạng thái và tình trạng. BorrowDetails có thể tham chiếu InstanceID để xác định cá thể cụ thể.")

    add_heading(doc, "4.2. Mô tả bảng", 2)
    add_table(doc, ["Bảng", "Khóa/quan hệ", "Ý nghĩa"], [
        ["Roles", "RoleID; RoleName duy nhất", "Danh mục vai trò."],
        ["Users", "UserID; FK RoleID; Username/UserCode duy nhất", "Tài khoản và hồ sơ người dùng."],
        ["Rooms", "RoomID; RoomCode/RoomName duy nhất", "Vị trí quản lý thiết bị."],
        ["DeviceCategories", "CategoryID; CategoryName duy nhất", "Phân loại thiết bị."],
        ["Devices", "DeviceID; FK CategoryID, RoomID; DeviceCode duy nhất", "Thiết bị tổng hợp và số lượng."],
        ["DeviceInstances", "InstanceID; FK DeviceID; AssetCode duy nhất", "Cá thể thiết bị."],
        ["BorrowTickets", "TicketID; FK UserID/ReturnedBy; TicketCode duy nhất", "Thông tin đầu phiếu và trạng thái."],
        ["BorrowDetails", "BorrowDetailID; FK TicketID/DeviceID/InstanceID", "Thiết bị/cá thể và số lượng mượn/trả."],
        ["ReturnRequests", "ReturnRequestID; TicketID duy nhất", "Yêu cầu trả đang chờ duyệt."],
        ["ReturnRequestDetails", "FK request/device/instance", "Từng dòng trả và tình trạng."],
        ["AppSettings", "SettingKey", "Cấu hình nghiệp vụ và thông tin đơn vị."],
    ], [1.55, 2.3, 2.3], font_size=9.6)

    add_heading(doc, "4.3. Ràng buộc và chỉ mục", 2)
    add_bullets(doc, [
        "CK_Devices_TotalQty và CK_Devices_AvailableQty ngăn số lượng âm hoặc số có sẵn vượt tổng.",
        "CK_BorrowDetails_Quantity yêu cầu Quantity > 0.",
        "CK_BorrowDetails_ReturnedQuantity bảo đảm 0 <= ReturnedQuantity <= Quantity.",
        "CK_BorrowTickets_Status giới hạn sáu trạng thái PENDING, BORROWING, RETURN_PENDING, RETURNED, REJECTED và CANCELLED.",
        "Unique index bảo vệ Username, UserCode, DeviceCode, AssetCode, RoomCode, RoomName, CategoryName và TicketCode.",
        "UX_ReturnRequests_TicketID bảo đảm một phiếu chỉ có tối đa một yêu cầu trả đang tồn tại.",
    ])

    add_heading(doc, "4.4. Schema thích nghi và migration", 2)
    add_para(doc, "DbSchemaHelper.HasColumn kiểm tra sự tồn tại của cột và cache kết quả theo DataSource, database, bảng và cột. Nhờ đó mã có thể làm việc với các phiên bản schema khác nhau, ví dụ ReturnDate/ActualReturnDate, có hoặc không có TicketCode, Purpose, ReturnNote hoặc InstanceID.")
    add_para(doc, "Mặt tích cực là giảm lỗi khi chạy trên database cũ. Mặt hạn chế là ứng dụng thực hiện DDL trong thời gian chạy: thêm cột, tạo bảng, thay đổi kiểu Note và thậm chí xóa/tạo lại check constraint. Cách này đòi hỏi tài khoản ứng dụng có quyền ALTER/CREATE, có thể gây khóa schema và khó kiểm soát phiên bản.")
    add_callout(doc, "Khuyến nghị", "Đưa toàn bộ thay đổi schema sang bộ migration có version, chạy một lần khi triển khai. Trong luồng nghiệp vụ chỉ kiểm tra phiên bản schema, không lặp lại ALTER TABLE hoặc DROP/ADD CONSTRAINT.")

    add_heading(doc, "4.5. Số liệu database tại thời điểm review", 2)
    add_table(doc, ["Bảng", "Số dòng"], [
        ["Users", "108"],
        ["Roles", "4"],
        ["Rooms", "24"],
        ["DeviceCategories", "12"],
        ["Devices", "38"],
        ["DeviceInstances", "1.050"],
        ["BorrowTickets", "40"],
        ["BorrowDetails", "66"],
        ["ReturnRequests", "4"],
        ["ReturnRequestDetails", "9"],
        ["AppSettings", "11"],
    ], [3.2, 2.95], font_size=10.5)

    add_heading(doc, "4.6. Phân tích dữ liệu người dùng", 2)
    add_para(doc, "Database hiện có 108 người dùng và 4 vai trò. Cơ cấu người dùng tập trung chủ yếu ở nhóm Sinh viên, phù hợp với bài toán mượn thiết bị trong môi trường đào tạo. Tại thời điểm kiểm tra, toàn bộ 108 tài khoản đang hoạt động và không có tài khoản bị khóa vĩnh viễn.")
    add_table(doc, ["Vai trò", "Số người", "Tỷ lệ", "Đang hoạt động", "Bị khóa"], [
        ["Sinh viên", "82", "75,9%", "82", "0"],
        ["Giảng viên", "19", "17,6%", "19", "0"],
        ["Admin", "5", "4,6%", "5", "0"],
        ["Thủ kho", "2", "1,9%", "2", "0"],
        ["Tổng", "108", "100%", "108", "0"],
    ], [1.35, 0.9, 0.85, 1.4, 1.1], font_size=10)
    add_para(doc, "Phân bố đơn vị tương đối đồng đều ở các nhóm chính: Công nghệ thông tin có 17 người; Môi trường, Kinh tế và Xây dựng mỗi nhóm 16 người; Điện - Điện tử có 13 người; Ngoại ngữ và Cơ khí mỗi nhóm 11 người. Ngoài ra còn một số nhãn đơn vị chỉ có một người.")
    add_callout(doc, "Vấn đề chuẩn hóa dữ liệu", "Các giá trị “Khoa học máy tinh” và “Khoa Học Máy Tính” có khả năng cùng chỉ một đơn vị nhưng khác chữ hoa và lỗi chính tả. Department hiện là chuỗi tự do, vì vậy nên tách thành bảng Departments và tham chiếu bằng khóa ngoại.")
    add_para(doc, "Kiểm tra dữ liệu bắt buộc cho thấy không có người dùng thiếu email hoặc số điện thoại. Username và UserCode được bảo vệ bởi unique index, giúp ngăn trùng tài khoản và mã người dùng ở tầng cơ sở dữ liệu.")

    add_heading(doc, "4.7. Phân tích dữ liệu thiết bị và cá thể", 2)
    add_para(doc, "Hệ thống đang quản lý 38 bản ghi thiết bị tổng hợp thuộc 12 loại, tương ứng 1.050 cá thể. Tổng AvailableQuantity của bảng Devices là 1.013, nghĩa là 37 đơn vị đang không được tính là có thể cho mượn. Các nhóm có số lượng lớn nhất là Bộ phát Wifi 125, Chuột 122, Bàn phím 114, Màn hình 104 và Router 94.")
    add_table(doc, ["Loại thiết bị", "Số mẫu", "Tổng cá thể", "Có thể mượn", "Chênh lệch"], [
        ["Bộ phát Wifi", "2", "125", "124", "1"],
        ["Chuột", "2", "122", "122", "0"],
        ["Bàn phím", "2", "114", "107", "7"],
        ["Màn hình", "2", "104", "104", "0"],
        ["Router", "2", "94", "92", "2"],
        ["Camera", "3", "90", "84", "6"],
        ["Laptop", "7", "79", "76", "3"],
        ["Loa", "3", "79", "77", "2"],
        ["Micro", "3", "78", "75", "3"],
        ["Máy tính để bàn", "3", "67", "65", "2"],
        ["Máy chiếu", "6", "52", "47", "5"],
        ["Thiết bị thí nghiệm", "3", "46", "40", "6"],
        ["Tổng", "38", "1.050", "1.013", "37"],
    ], [1.75, 0.75, 1.0, 1.1, 1.0], font_size=9.2)
    add_para(doc, "Theo bảng DeviceInstances, 1.042 cá thể có trạng thái Có sẵn, 7 cá thể Bảo trì và 1 cá thể Mất. Con số 1.042 cá thể Có sẵn cao hơn AvailableQuantity tổng hợp 1.013 là dấu hiệu dữ liệu lịch sử chưa đồng bộ ở mức cá thể. Cụ thể, các phiếu đang hoạt động đã trừ số lượng ở Devices nhưng 30 dòng BorrowDetails không có InstanceID nên không thể chuyển đúng DeviceInstance sang trạng thái Đang mượn.")
    add_table(doc, ["Trạng thái cá thể", "Số lượng", "Tỷ lệ"], [
        ["Có sẵn", "1.042", "99,24%"],
        ["Bảo trì", "7", "0,67%"],
        ["Mất", "1", "0,09%"],
        ["Tổng", "1.050", "100%"],
    ], [2.6, 1.5, 2.05], font_size=10.2)
    add_para(doc, "Trường Condition chưa được khai báo rõ cho 905/1.050 cá thể, tương đương khoảng 86,2%. Mã nguồn đang coi NULL hoặc chuỗi rỗng là “Tốt” trong nhiều truy vấn. Cách tương thích này giúp chương trình hoạt động với dữ liệu cũ nhưng làm giảm chất lượng phân tích tình trạng thực tế.")
    add_callout(doc, "Đề xuất làm sạch", "Chuẩn hóa Condition về tập giá trị Tốt, Bảo trì, Hỏng, Mất, Trầy xước, Thiếu phụ kiện; cập nhật dữ liệu NULL thành Tốt sau khi kiểm kê; thêm CHECK constraint hoặc bảng danh mục tình trạng.")

    add_heading(doc, "4.8. Phân bố thiết bị theo phòng", 2)
    add_para(doc, "Toàn bộ 38 thiết bị đều đã được gán RoomID, không có thiết bị thiếu vị trí. Thiết bị tập trung nhiều nhất ở Phòng B301 với 126 cá thể, B202 với 117, B201 với 109, B102 với 101 và B101 với 94. Nhóm năm phòng này chiếm 547/1.050 cá thể, tương đương khoảng 52,1% tổng kho.")
    add_table(doc, ["Phòng", "Số mẫu thiết bị", "Tổng cá thể"], [
        ["B301", "2", "126"],
        ["B202", "2", "117"],
        ["B201", "2", "109"],
        ["B102", "2", "101"],
        ["B101", "2", "94"],
        ["A302", "2", "83"],
        ["A301", "2", "75"],
        ["A202", "2", "67"],
        ["A201", "3", "64"],
        ["A102", "3", "62"],
        ["Các phòng còn lại", "14", "152"],
    ], [2.2, 1.75, 2.2], font_size=10)
    add_para(doc, "Việc tập trung tài sản ở một số phòng giúp ưu tiên kiểm kê và bảo trì theo rủi ro. Tuy nhiên, hiện hệ thống gắn Devices với một RoomID cố định trong khi phiếu mượn lưu “Phòng sử dụng” trong chuỗi Note. Nếu thiết bị có thể di chuyển giữa các phòng, nên bổ sung bảng DeviceMovements hoặc lưu RoomID trực tiếp trên BorrowTickets.")

    add_heading(doc, "4.9. Phân tích dữ liệu phiếu mượn và trả", 2)
    add_para(doc, "Có 40 phiếu phát sinh trong ba tháng từ tháng 4 đến tháng 6 năm 2026. Trong đó 28 phiếu đã trả, 10 phiếu đang mượn và 2 phiếu bị từ chối. Không có phiếu đang hoạt động quá hạn tại thời điểm kiểm tra. Bốn bản ghi ReturnRequests cho thấy đang có dữ liệu yêu cầu trả được lưu có cấu trúc, dù trạng thái tổng hợp của phiếu hiện không có RETURN_PENDING.")
    add_table(doc, ["Trạng thái", "Số phiếu", "Tỷ lệ", "Tổng SL trên phiếu"], [
        ["RETURNED", "28", "70%", "29"],
        ["BORROWING", "10", "25%", "30"],
        ["REJECTED", "2", "5%", "7"],
        ["Tổng", "40", "100%", "66"],
    ], [1.6, 1.0, 0.9, 2.65], font_size=10)
    add_table(doc, ["Tháng", "Số phiếu", "Số lượng thiết bị"], [
        ["04/2026", "9", "10"],
        ["05/2026", "15", "19"],
        ["06/2026", "16", "37"],
    ], [2.2, 1.8, 2.15], font_size=10.5)
    add_para(doc, "Số phiếu tăng từ 9 lên 16 trong ba tháng, còn số thiết bị trên phiếu tăng từ 10 lên 37. Dữ liệu này cho thấy mức sử dụng hệ thống đang tăng, nhưng khoảng thời gian quan sát còn ngắn nên chưa đủ để kết luận xu hướng dài hạn hoặc mùa vụ.")

    add_heading(doc, "4.10. Thiết bị được mượn nhiều", 2)
    add_table(doc, ["Hạng", "Mã", "Tên thiết bị", "Số lượt", "Số phiếu"], [
        ["1", "TB09", "Bàn phím Logitech K120", "23", "15"],
        ["2", "TB33", "Bàn phím Dell KB216", "6", "3"],
        ["3", "TB24", "Camera Logitech C920", "4", "4"],
        ["4", "TB37", "Bộ dụng cụ Arduino Uno", "3", "3"],
        ["5", "TB12", "Bộ thí nghiệm Arduino Starter Kit", "3", "2"],
        ["6", "TB02", "Laptop Acer Aspire 5 A515", "3", "3"],
        ["7", "TB25", "Camera Hikvision DS-2CD", "2", "2"],
        ["8", "TB06", "Camera Hikvision DS-2CE16D0T", "2", "2"],
        ["9", "TB14", "Laptop Dell Latitude 5420", "2", "2"],
        ["10", "TB22", "Máy chiếu BenQ MW560", "2", "2"],
    ], [0.55, 0.75, 2.85, 0.9, 1.1], font_size=9.2)
    add_para(doc, "TB09 chiếm 23/66 đơn vị đã ghi nhận trong BorrowDetails, tương đương khoảng 34,8%, và xuất hiện trong 15/40 phiếu. Đây là mức tập trung sử dụng cao, cần ưu tiên kiểm tra hao mòn, chuẩn bị thiết bị thay thế và đánh giá có cần tăng số lượng bàn phím cùng loại hay không.")

    add_heading(doc, "4.11. Đánh giá chất lượng dữ liệu tổng thể", 2)
    add_table(doc, ["Tiêu chí", "Kết quả", "Đánh giá"], [
        ["Khóa ngoại mồ côi", "Không phát hiện", "Tốt"],
        ["Mã thiết bị/cá thể trùng", "Không phát hiện", "Tốt"],
        ["Số lượng âm hoặc vượt tổng", "Không phát hiện", "Tốt"],
        ["Thiết bị thiếu mã/phòng", "0", "Tốt"],
        ["Người dùng thiếu email/điện thoại", "0", "Tốt"],
        ["BorrowDetails hoạt động thiếu InstanceID", "30 dòng", "Cần xử lý ưu tiên"],
        ["Cá thể thiếu Condition chuẩn", "905 dòng", "Cần làm sạch"],
        ["Tên đơn vị chưa chuẩn hóa", "Có biến thể chữ hoa/chính tả", "Cần danh mục hóa"],
        ["ReturnRequests còn tồn tại nhưng không có RETURN_PENDING", "4 request", "Cần đối soát trạng thái"],
    ], [2.6, 1.7, 1.85], font_size=9.7)
    add_para(doc, "Nhìn chung, dữ liệu có nền tảng quan hệ tốt: khóa ngoại, unique index và check constraint đang hoạt động; không có bản ghi mồ côi hoặc số lượng bất hợp lệ. Các vấn đề chính xuất phát từ quá trình nâng cấp từ quản lý số lượng tổng hợp sang quản lý từng cá thể và từ các trường nhập tự do chưa được chuẩn hóa.")


def chapter_implementation(doc):
    add_heading(doc, "CHƯƠNG 5. HIỆN THỰC CÁC CHỨC NĂNG", 1)
    add_heading(doc, "5.1. Đăng nhập và bảo mật mật khẩu", 2)
    add_para(doc, "PasswordHelper dùng PasswordHasher<object> của Microsoft.Extensions.Identity.Core. Đây là lựa chọn tốt hơn lưu mật khẩu rõ hoặc tự băm SHA đơn giản vì format của Identity bao gồm salt, tham số thuật toán và hỗ trợ nâng cấp.")
    add_para(doc, "AuthRepository theo dõi FailedLoginCount và LockoutUntil. Sau 5 lần sai, tài khoản bị tạm khóa 5 phút. Username không tồn tại cũng bị giới hạn bằng ConcurrentDictionary trong bộ nhớ để tránh cho phép thử vô hạn trong cùng phiên ứng dụng. Thông báo đăng nhập không tiết lộ tài khoản có tồn tại hay không.")

    add_heading(doc, "5.2. Quản lý thiết bị và cá thể", 2)
    add_para(doc, "Màn hình thiết bị hỗ trợ phân trang 20 dòng, tìm theo tên/ghi chú, lọc loại và trạng thái. SaveDevice kiểm tra mã trùng, phòng tồn tại và không cho giảm tổng số lượng xuống dưới số đang mượn. Xóa thiết bị là xóa mềm: đặt trạng thái Ngừng sử dụng, số có sẵn bằng 0 và chuyển các cá thể đang Có sẵn sang Ngừng sử dụng.")
    add_para(doc, "DeviceInstanceRepository tự sinh AssetCode theo mã thiết bị và số thứ tự nhỏ nhất chưa dùng. Khi thêm, sửa hoặc xóa cá thể, repository cập nhật lại TotalQuantity và AvailableQuantity của thiết bị cha trong cùng transaction.")

    add_heading(doc, "5.3. Quản lý người dùng", 2)
    add_para(doc, "Màn hình người dùng có phân trang phía server, tìm kiếm theo họ tên/mã/email/username, lọc vai trò và trạng thái. Mật khẩu mới phải có ít nhất 6 ký tự, gồm chữ và số, sau đó được băm trước khi lưu. Xóa người dùng là đặt IsActive = 0 và IsLocked = 1 để giữ nguyên lịch sử phiếu.")
    add_para(doc, "Chức năng xuất Excel dùng EPPlus. License được đặt ở chế độ NonCommercialPersonal, phù hợp bài tập học tập nhưng cần xem xét lại giấy phép nếu triển khai thương mại.")

    add_heading(doc, "5.4. Lập và phê duyệt phiếu mượn", 2)
    add_para(doc, "UcLapthietbi chỉ đưa các thiết bị có ít nhất một cá thể Có sẵn, tình trạng Tốt và phòng chưa ngừng sử dụng vào danh sách. Người dùng chọn cá thể cụ thể; mỗi dòng có số lượng 1. Cấu hình CaiDat_ToiDaThietBi giới hạn tổng cá thể trên phiếu.")
    add_para(doc, "CreatePendingTicket tạo đầu phiếu và chi tiết trong transaction. Nếu cấu hình yêu cầu phê duyệt, phiếu ở PENDING. Nếu không, giao diện gọi BorrowApprovalService ngay. Khi duyệt, dịch vụ đọc Devices và DeviceInstances với UPDLOCK, HOLDLOCK, kiểm tra lại trạng thái/tồn, trừ AvailableQuantity và đổi trạng thái cá thể sang Đang mượn.")

    add_heading(doc, "5.5. Trả đủ, trả thiếu và báo lỗi", 2)
    add_para(doc, "UcTrathietbi cho phép chọn phiếu đang mượn, nhập số lượng trả và tình trạng. ReturnIssueDialog hỗ trợ các tình trạng Tốt, Hỏng hóc, Bảo trì, Mất, Trầy xước, Thiếu phụ kiện và Khác. Nếu trả thiếu hoặc tình trạng không tốt, người dùng bắt buộc nhập cảnh báo/ghi chú.")
    add_para(doc, "Yêu cầu trả được lưu vào ReturnRequests/ReturnRequestDetails và đồng thời có payload dự phòng trong BorrowTickets.ReturnNote. Khi duyệt, dữ liệu được cập nhật trong transaction. Cá thể tốt trở về Có sẵn; cá thể lỗi chuyển Bảo trì; ReturnedQuantity tăng; trạng thái phiếu phụ thuộc số lượng còn mượn.")

    add_heading(doc, "5.6. Dashboard và báo cáo", 2)
    add_para(doc, "Dashboard tải snapshot ở background thread bằng Task.Run để giảm treo giao diện. Các chỉ số gồm tổng thiết bị, tổng số đang mượn, phiếu quá hạn, thiết bị có vấn đề, phiếu trả hôm nay, danh sách đang mượn có phân trang và thống kê 6 tháng.")
    add_para(doc, "Màn hình báo cáo có ba nhóm: tổng hợp mượn trả theo tháng, top 10 thiết bị được mượn nhiều và phiếu quá hạn. Dữ liệu có thể xuất thành nhiều sheet Excel. Chức năng hiện xuất giá trị dạng chuỗi; phiên bản sau nên xuất đúng kiểu ngày/số, định dạng bảng và thêm biểu đồ.")

    add_heading(doc, "5.7. Xóa mềm và thùng rác", 2)
    add_para(doc, "Thiết bị, cá thể, phòng và người dùng đều có chiến lược xóa mềm. RecycleBinRepository tổng hợp các bản ghi đã ngừng sử dụng, cho phép khôi phục hoặc xóa vĩnh viễn. Trước khi xóa vĩnh viễn, hệ thống kiểm tra lịch sử BorrowDetails/BorrowTickets và quan hệ thiết bị - phòng để tránh phá dữ liệu tham chiếu.")

    add_heading(doc, "5.8. Logging và xử lý lỗi", 2)
    add_para(doc, "Serilog ghi file logs/app-yyyyMMdd.log ở mức Debug trở lên. Program bắt lỗi UI thread và AppDomain, ghi log rồi hiển thị thông báo. Nhiều màn hình cũng ghi lỗi tải dữ liệu không nghiêm trọng. Đây là nền tảng tốt cho truy vết, nhưng log lịch sử nghiệp vụ hiện chủ yếu là chuỗi văn bản; bảng ActivityLogs đang có 0 dòng và chưa được tích hợp.")


def chapter_review(doc):
    add_heading(doc, "CHƯƠNG 6. REVIEW MÃ NGUỒN VÀ ĐÁNH GIÁ KỸ THUẬT", 1)
    add_heading(doc, "6.1. Phạm vi và số liệu review", 2)
    add_table(doc, ["Chỉ số", "Giá trị"], [
        ["Tệp C# được thống kê", "60"],
        ["Số dòng C# ước tính", "15.240"],
        ["Tệp lớn nhất", "UcDanhsachphieu.cs - 1.040 dòng"],
        ["Tệp nghiệp vụ lớn", "ReturnApprovalService.cs - 473 dòng"],
        ["Cấu hình build", "Release, net8.0-windows"],
        ["Kết quả build", "Thành công, 0 lỗi, 56 cảnh báo"],
        ["Test project", "Không tìm thấy xUnit/NUnit/MSTest"],
    ], [2.4, 3.75], font_size=10.5)

    add_heading(doc, "6.2. Điểm mạnh", 2)
    add_bullets(doc, [
        "Nghiệp vụ mượn/trả có transaction và kiểm tra trạng thái trước khi cập nhật.",
        "Duyệt mượn dùng UPDLOCK/HOLDLOCK để giảm tranh chấp tồn kho.",
        "Mật khẩu sử dụng PasswordHasher chuẩn thay vì lưu rõ.",
        "Có tạm khóa đăng nhập và thông báo không làm lộ username hợp lệ.",
        "Phân quyền được gom trong PermissionSet, dễ quan sát.",
        "Truy vấn phần lớn dùng tham số, giảm nguy cơ SQL injection.",
        "Có phân trang phía server ở các danh sách lớn.",
        "Xóa mềm và khôi phục giữ được lịch sử nghiệp vụ.",
        "Có check constraint, unique index và foreign key ở database.",
        "Cơ chế schema detection hỗ trợ dữ liệu cũ và nhiều phiên bản cột.",
    ])

    add_heading(doc, "6.3. Phát hiện mức cao", 2)
    add_table(doc, ["Mã", "Phát hiện", "Tác động", "Khuyến nghị"], [
        ["H-01", "DDL/migration được chạy trong luồng tạo và duyệt phiếu; check constraint bị drop/add lặp lại.", "Yêu cầu quyền ALTER, khóa schema, tăng thời gian và rủi ro lỗi khi nhiều người dùng.", "Dùng migration có version; luồng nghiệp vụ chỉ kiểm tra version."],
        ["H-02", "Dữ liệu đang hoạt động có 30 BorrowDetails không gắn InstanceID.", "Số lượng tổng hợp giảm nhưng trạng thái từng cá thể vẫn Có sẵn; có thể chọn lại cá thể không đúng thực tế.", "Viết migration đối soát và gán cá thể; chặn phiếu mới thiếu InstanceID."],
        ["H-03", "Chưa có test tự động cho transaction mượn/trả và dữ liệu cạnh tranh.", "Hồi quy dễ làm sai tồn kho hoặc trạng thái phiếu.", "Bổ sung unit test helper/service và integration test trên database tạm."],
    ], [0.55, 2.25, 1.65, 1.7], font_size=9.2, header_fill="F4CCCC")

    add_heading(doc, "6.4. Phát hiện mức trung bình", 2)
    add_table(doc, ["Mã", "Phát hiện", "Tác động/chi tiết"], [
        ["M-01", "56 cảnh báo build.", "Chủ yếu CS8622 do chữ ký event handler, ba CS8602 trong ReturnIssueDialog, một CS8604 và bốn field designer không dùng."],
        ["M-02", "Các lớp GUI quá lớn.", "UcDanhsachphieu 1.040 dòng, UcTrathietbi 692 dòng, FrmMain 584 dòng; khó test và bảo trì."],
        ["M-03", "GUI gọi trực tiếp repository.", "Làm mờ ranh giới BLL/DAL, khó mock và khó kiểm thử độc lập."],
        ["M-04", "ReturnApprovalService dài và nhiều nhánh.", "Duyệt toàn phần, một phần, bulk/cá thể và tương thích schema cùng nằm trong một phương thức."],
        ["M-05", "Payload ReturnNote bị giới hạn 255 ký tự.", "BuildPayload dừng khi đủ 255 ký tự; nếu bảng ReturnRequests mất hoặc không tồn tại có thể thiếu dòng trả trong fallback."],
        ["M-06", "Dashboard ReturnedToday chỉ kiểm tra ReturnDate.", "Không dùng helper hỗ trợ ActualReturnDate như các repository khác."],
        ["M-07", "Role mapping dựa trên Contains.", "Tên vai trò mới có thể bị ánh xạ sai; quyền không được lưu thành ma trận rõ trong database."],
        ["M-08", "Nhiều AddWithValue.", "SQL Server có thể suy luận kiểu/độ dài không tối ưu và ảnh hưởng plan cache."],
    ], [0.65, 2.3, 3.2], font_size=9.5, header_fill="FCE5CD")

    add_heading(doc, "6.5. Phát hiện mức thấp và chất lượng mã", 2)
    add_table(doc, ["Mã", "Phát hiện", "Hướng xử lý"], [
        ["L-01", "Tên control như cmbNgayLap được dùng làm danh sách thiết bị.", "Đổi tên theo ý nghĩa để giảm nhầm lẫn."],
        ["L-02", "Biên bản trả dùng văn bản ASCII không dấu.", "Tạo mẫu DOCX/PDF tiếng Việt và lưu lịch sử."],
        ["L-03", "README chỉ có một dòng.", "Bổ sung hướng dẫn cài database, tài khoản mẫu, build và migration."],
        ["L-04", "Notebook schema cũ không phản ánh toàn bộ schema hiện tại.", "Xuất migration hoặc script SQL chuẩn trong source control."],
        ["L-05", "ActivityLogs chưa được sử dụng.", "Ghi audit log nghiệp vụ có cấu trúc hoặc loại bỏ bảng."],
        ["L-06", "Một số catch bỏ qua lỗi.", "Ghi log tối thiểu và phân biệt lỗi có thể phục hồi."],
    ], [0.65, 2.55, 2.95], font_size=9.7)

    add_heading(doc, "6.6. Đánh giá theo nguyên tắc thiết kế", 2)
    add_table(doc, ["Nguyên tắc", "Mức đáp ứng", "Nhận xét"], [
        ["Single Responsibility", "Trung bình", "Repository tương đối rõ; các UserControl và ReturnApprovalService còn quá nhiều trách nhiệm."],
        ["Separation of Concerns", "Trung bình", "Có thư mục phân lớp nhưng GUI vẫn gọi DAL trực tiếp."],
        ["Dependency Inversion", "Thấp", "Hầu hết lớp static, chưa có interface/DI nên khó mock."],
        ["Defensive programming", "Khá", "Kiểm tra trạng thái, số lượng, transaction và ràng buộc database."],
        ["Observability", "Trung bình", "Có file log nhưng thiếu correlation ID và audit log có cấu trúc."],
        ["Testability", "Thấp", "Static repository, MessageBox trong validation, không có test project."],
    ], [1.6, 1.2, 3.35], font_size=10)


def chapter_testing(doc):
    add_heading(doc, "CHƯƠNG 7. KIỂM THỬ VÀ XÁC MINH", 1)
    add_heading(doc, "7.1. Kết quả build", 2)
    add_code(doc, """Lệnh: dotnet build FrmProject.slnx -c Release
Kết quả: Build succeeded
Cảnh báo: 56
Lỗi: 0
Ngày xác minh: 05/06/2026""")
    add_para(doc, "Các cảnh báo không ngăn chương trình biên dịch nhưng cần được xử lý trước khi coi build là sạch. Dự án đã bật Nullable nhưng chưa bật TreatWarningsAsErrors, vì vậy lỗi nullable tiềm ẩn chưa bị chặn ở CI.")

    add_heading(doc, "7.2. Kiểm tra cơ sở dữ liệu chỉ đọc", 2)
    add_table(doc, ["Kiểm tra", "Kết quả"], [
        ["TotalQuantity/AvailableQuantity âm hoặc vượt giới hạn", "0 lỗi"],
        ["ReturnedQuantity âm hoặc lớn hơn Quantity", "0 lỗi"],
        ["BorrowDetails mồ côi theo TicketID", "0 lỗi"],
        ["BorrowDetails mồ côi theo DeviceID", "0 lỗi"],
        ["DeviceCode trùng", "0 nhóm trùng"],
        ["AssetCode trùng", "0 nhóm trùng"],
        ["Cá thể Đang mượn không có phiếu hoạt động", "0 lỗi"],
        ["Chi tiết phiếu đang hoạt động thiếu InstanceID", "30 dòng"],
    ], [4.35, 1.8], font_size=10.3)
    add_para(doc, "Dữ liệu lịch sử thiếu InstanceID là vấn đề tương thích giữa schema cũ và thiết kế cá thể mới. Các phiếu này vẫn trừ AvailableQuantity của Devices, nhưng không đổi trạng thái một DeviceInstance cụ thể. Vì vậy dữ liệu tổng và dữ liệu cá thể có thể cho hai cách hiểu khác nhau.")

    add_heading(doc, "7.3. Bộ test case nghiệm thu đề xuất", 2)
    tests = [
        ["TC-01", "Đăng nhập đúng", "Tài khoản hoạt động, mật khẩu đúng", "Mở FrmMain đúng vai trò."],
        ["TC-02", "Khóa đăng nhập", "Sai 5 lần", "Tạm khóa 5 phút, còn số lần thử được hiển thị."],
        ["TC-03", "Tài khoản ngừng/khóa", "Mật khẩu đúng", "Không cho truy cập."],
        ["TC-04", "Thêm thiết bị", "Mã mới, phòng hợp lệ", "Tạo Devices, số lượng đúng."],
        ["TC-05", "Mã thiết bị trùng", "DeviceCode đã có", "Từ chối lưu."],
        ["TC-06", "Giảm số lượng dưới số đang mượn", "Total mới quá nhỏ", "Từ chối cập nhật."],
        ["TC-07", "Thêm cá thể", "AssetCode mới", "Tổng và số có sẵn đồng bộ."],
        ["TC-08", "Lập phiếu hợp lệ", "Các cá thể Tốt/Có sẵn", "Tạo PENDING cùng chi tiết."],
        ["TC-09", "Trùng cá thể trong phiếu", "Chọn lại AssetCode", "Không thêm trùng."],
        ["TC-10", "Vượt giới hạn cấu hình", "Tổng > CaiDat_ToiDaThietBi", "Hiển thị cảnh báo."],
        ["TC-11", "Duyệt mượn", "Phiếu PENDING đủ tồn", "BORROWING, giảm tồn, cá thể Đang mượn."],
        ["TC-12", "Duyệt mượn cạnh tranh", "Hai phiên cùng cá thể", "Chỉ một phiên thành công."],
        ["TC-13", "Từ chối mượn", "Phiếu PENDING", "REJECTED và ghi lý do."],
        ["TC-14", "Trả đủ tốt", "Tất cả dòng trả đủ", "RETURNED, tăng tồn, cá thể Có sẵn."],
        ["TC-15", "Trả hỏng", "Tình trạng Hỏng hóc", "Cá thể Bảo trì; không tăng số có sẵn."],
        ["TC-16", "Trả một phần", "Chỉ chọn một số dòng", "Phần còn lại RETURN_PENDING/BORROWING phù hợp."],
        ["TC-17", "Từ chối trả", "Phiếu RETURN_PENDING", "Quay lại BORROWING."],
        ["TC-18", "Xóa mềm thiết bị", "Thiết bị có lịch sử", "Ngừng sử dụng, vẫn giữ lịch sử."],
        ["TC-19", "Xóa vĩnh viễn có tham chiếu", "Còn BorrowDetails", "Từ chối xóa."],
        ["TC-20", "Xuất Excel", "Báo cáo có dữ liệu", "Tệp XLSX mở được, đúng sheet."],
    ]
    add_table(doc, ["Mã", "Kịch bản", "Dữ liệu/Thao tác", "Kết quả mong đợi"], tests, [0.65, 1.55, 2.0, 1.95], font_size=8.8)

    add_heading(doc, "7.4. Chiến lược test tự động đề xuất", 2)
    add_bullets(doc, [
        "Unit test PasswordHelper, ValidationHelper, ReturnRequestHelper và ánh xạ vai trò.",
        "Unit test state transition bằng service không phụ thuộc MessageBox/WinForms.",
        "Integration test repository trên database riêng, chạy migration trước mỗi nhóm test.",
        "Concurrency test duyệt cùng một cá thể từ hai transaction.",
        "Regression test trả một phần nhiều lần, trả hỏng, thiết bị cha ngừng sử dụng.",
        "Smoke test khởi động ứng dụng với database schema cũ và schema mới.",
        "CI chạy restore, build, test và fail khi có warning mới.",
    ])


def chapter_deployment(doc):
    add_heading(doc, "CHƯƠNG 8. TRIỂN KHAI VÀ HƯỚNG DẪN SỬ DỤNG", 1)
    add_heading(doc, "8.1. Yêu cầu môi trường", 2)
    add_bullets(doc, [
        "Windows 10 hoặc Windows 11.",
        ".NET 8 Desktop Runtime hoặc bản publish self-contained.",
        "SQL Server Express, instance mặc định localhost\\SQLEXPRESS.",
        "Database QuanLyThietBi và tài khoản Windows có quyền kết nối.",
        "Quyền ghi thư mục logs cạnh tệp thực thi.",
        "Microsoft Excel không bắt buộc; tệp XLSX được tạo trực tiếp bằng EPPlus.",
    ])

    add_heading(doc, "8.2. Cấu hình kết nối", 2)
    add_code(doc, """<connectionStrings>
  <add name="QuanLyThietBi"
       connectionString="Server=localhost\\SQLEXPRESS;
                         Database=QuanLyThietBi;
                         Trusted_Connection=True;
                         TrustServerCertificate=True;"
       providerName="Microsoft.Data.SqlClient" />
</connectionStrings>""")
    add_para(doc, "Chuỗi kết nối dùng Windows Authentication và TrustServerCertificate. Trong môi trường nhiều máy, cần thay DataSource, cấu hình firewall/TCP và quản lý quyền database theo nguyên tắc tối thiểu.")

    add_heading(doc, "8.3. Quy trình sử dụng", 2)
    add_numbered(doc, [
        "Đăng nhập bằng tài khoản đã được cấp; xử lý cảnh báo nếu tài khoản bị khóa.",
        "Quản trị viên cấu hình thông tin đơn vị, thời hạn mượn, số thiết bị tối đa và yêu cầu phê duyệt.",
        "Thủ kho cập nhật phòng, loại thiết bị, thiết bị và mã cá thể.",
        "Người dùng hoặc thủ kho lập phiếu, chọn đúng cá thể và hạn trả.",
        "Thủ kho/Admin duyệt phiếu trong danh sách phiếu; hệ thống chuyển cá thể sang Đang mượn.",
        "Khi trả, người dùng chọn phiếu, nhập tình trạng và gửi yêu cầu.",
        "Thủ kho/Admin kiểm tra và duyệt/từ chối yêu cầu trả.",
        "Theo dõi dashboard, báo cáo quá hạn và xuất Excel khi cần.",
        "Dùng thùng rác để khôi phục dữ liệu xóa nhầm; chỉ xóa vĩnh viễn khi không còn lịch sử tham chiếu.",
    ])

    add_heading(doc, "8.4. Khuyến nghị vận hành", 2)
    add_bullets(doc, [
        "Không sửa trực tiếp trạng thái và số lượng trong SQL Server nếu không có biên bản đối soát.",
        "Sao lưu database trước khi nâng cấp schema hoặc xóa vĩnh viễn.",
        "Định kỳ kiểm tra phiếu BORROWING/RETURN_PENDING và dữ liệu InstanceID.",
        "Giới hạn quyền ALTER/CREATE sau khi chuyển migration ra khỏi ứng dụng.",
        "Bảo vệ thư mục log vì có tên người dùng và thông tin lỗi.",
        "Kiểm tra license EPPlus trước khi sử dụng ngoài mục đích phi thương mại.",
    ])


def chapter_future(doc):
    add_heading(doc, "CHƯƠNG 9. ĐÁNH GIÁ VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_heading(doc, "9.1. Kết quả đạt được", 2)
    add_bullets(doc, [
        "Hoàn thành quy trình nghiệp vụ mượn trả nhiều trạng thái, có phê duyệt.",
        "Quản lý thiết bị ở cả mức tổng hợp và mức cá thể.",
        "Có phân quyền ba mức, bảo mật mật khẩu và giới hạn đăng nhập sai.",
        "Có transaction, locking, ràng buộc database, xóa mềm và thùng rác.",
        "Có dashboard, báo cáo và xuất Excel.",
        "Có cấu hình nghiệp vụ, logging và xử lý lỗi toàn cục.",
        "Build Release thành công trên .NET 8.",
    ])

    add_heading(doc, "9.2. Hạn chế", 2)
    add_table(doc, ["Hạn chế", "Ảnh hưởng"], [
        ["Không có test tự động", "Khó đảm bảo an toàn khi sửa nghiệp vụ tồn kho."],
        ["Migration chạy trong ứng dụng", "Đòi quyền cao, khó triển khai và có thể khóa schema."],
        ["GUI/DAL liên kết chặt", "Khó thay database, khó mock và khó test."],
        ["Dữ liệu cũ thiếu InstanceID", "Tổng kho và trạng thái cá thể chưa đồng nhất."],
        ["Chưa có email nhắc hạn", "Cảnh báo chỉ hiển thị khi mở ứng dụng."],
        ["Ứng dụng desktop đồng bộ", "Khả năng mở rộng nhiều người dùng và truy cập từ xa hạn chế."],
        ["Audit log chưa có cấu trúc", "Khó truy vết ai thay đổi dữ liệu nào."],
    ], [2.6, 3.55], font_size=10.2)

    add_heading(doc, "9.3. Lộ trình phát triển", 2)
    add_table(doc, ["Giai đoạn", "Mục tiêu", "Công việc"], [
        ["1 - Ổn định", "Giảm lỗi hồi quy", "Sửa 56 warning; thêm migration version; test helper và service; chuẩn hóa README."],
        ["2 - Làm sạch dữ liệu", "Đồng bộ kho/cá thể", "Đối soát 30 dòng thiếu InstanceID; thêm constraint/quy tắc chặn dữ liệu mới."],
        ["3 - Tái cấu trúc", "Tăng khả năng bảo trì", "Tách service/interface, DI, DTO, giảm UserControl lớn, bỏ SQL khỏi GUI."],
        ["4 - Vận hành", "Tăng truy vết", "Audit log database, correlation ID, health check, backup/restore có hướng dẫn."],
        ["5 - Mở rộng", "Tăng trải nghiệm", "QR/barcode, email nhắc hạn, in PDF, biểu đồ, import/export chuẩn."],
        ["6 - Hiện đại hóa", "Đa người dùng", "ASP.NET Core API + web/mobile hoặc WPF client; xác thực tập trung."],
    ], [1.1, 1.6, 3.45], font_size=9.8)

    add_heading(doc, "9.4. Kiến trúc mục tiêu đề xuất", 2)
    add_para(doc, "Phiên bản tiếp theo nên tách thành Presentation, Application, Domain và Infrastructure. Application định nghĩa use case và interface repository; Infrastructure hiện thực SQL Server; WinForms chỉ binding dữ liệu và gọi use case. DbContext hoặc Dapper có thể được cân nhắc, nhưng không bắt buộc thay ADO.NET nếu repository được chuẩn hóa.")
    add_para(doc, "Các trạng thái nên dùng enum/domain type ở C# và bảng/constraint ổn định ở database. Migration cần được quản lý bằng script version hoặc công cụ như DbUp/FluentMigrator/EF Core Migrations. Việc thay đổi kiến trúc nên thực hiện từng module, bắt đầu từ Auth và Borrow/Return vì có rủi ro cao nhất.")


def conclusion_appendices(doc):
    add_heading(doc, "PHẦN III. KẾT LUẬN", 1)
    add_para(doc, "Hệ thống quản lý mượn trả thiết bị đã vượt phạm vi của một chương trình CRUD cơ bản. Sản phẩm có quản lý cá thể, vòng đời phiếu, phê duyệt, trả một phần, theo dõi tình trạng, phân quyền, khóa đăng nhập, dashboard, báo cáo, xóa mềm và logging. Việc sử dụng transaction và khóa cập nhật ở các luồng quan trọng cho thấy dự án đã chú ý đến tính đúng đắn của dữ liệu.")
    add_para(doc, "Qua review, điểm cần ưu tiên không phải bổ sung thêm nhiều màn hình mà là tăng độ tin cậy: chuẩn hóa migration, xử lý dữ liệu lịch sử thiếu InstanceID, bổ sung test tự động, sửa cảnh báo nullable và tách các lớp lớn. Khi hoàn thành các bước này, dự án có thể trở thành nền tảng tốt để mở rộng sang hệ thống nhiều người dùng, API và giao diện hiện đại.")

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "[1] Microsoft. .NET 8 Documentation. https://learn.microsoft.com/dotnet/",
        "[2] Microsoft. Windows Forms Documentation. https://learn.microsoft.com/dotnet/desktop/winforms/",
        "[3] Microsoft. Microsoft.Data.SqlClient Documentation. https://learn.microsoft.com/sql/connect/ado-net/",
        "[4] Microsoft. ASP.NET Core Identity PasswordHasher. https://learn.microsoft.com/dotnet/api/microsoft.aspnetcore.identity.passwordhasher-1",
        "[5] Microsoft. SQL Server Transaction Locking and Row Versioning Guide.",
        "[6] Serilog Documentation. https://serilog.net/",
        "[7] EPPlus Documentation. https://epplussoftware.com/",
        "[8] Robert C. Martin. Clean Architecture. Prentice Hall.",
        "[9] Martin Fowler. Patterns of Enterprise Application Architecture. Addison-Wesley.",
    ]
    for ref in refs:
        add_para(doc, ref, first_indent=False)

    add_heading(doc, "PHỤ LỤC A. CẤU TRÚC MÃ NGUỒN", 1)
    add_code(doc, """FrmProject/
├─ BLL/
│  ├─ BorrowApprovalService.cs
│  ├─ ReturnApprovalService.cs
│  └─ ReturnRequestHelper.cs
├─ DAL/
│  ├─ DbHelper.cs, DbSchemaHelper.cs
│  ├─ AuthRepository.cs, UserRepository.cs
│  ├─ DeviceRepository.cs, DeviceInstanceRepository.cs
│  ├─ BorrowTicketRepository.cs, ReturnTicketRepository.cs
│  ├─ ReturnRequestRepository.cs, TicketListRepository.cs
│  └─ DashboardRepository.cs, ReportRepository.cs, ...
├─ GUI/
│  ├─ FrmLogin.cs, FrmMain.cs
│  └─ UcThietbi.cs, UcLapthietbi.cs, UcTrathietbi.cs, ...
├─ Models/
├─ DataStructures/
├─ Data/
├─ Program.cs
└─ FrmProject.csproj""")

    add_heading(doc, "PHỤ LỤC B. CẤU HÌNH DỰ ÁN", 1)
    add_table(doc, ["Thành phần", "Giá trị"], [
        ["TargetFramework", "net8.0-windows"],
        ["UI", "Windows Forms"],
        ["Nullable", "enable"],
        ["Database client", "Microsoft.Data.SqlClient 6.1.4"],
        ["Excel", "EPPlus 8.5.3"],
        ["Password hashing", "Microsoft.Extensions.Identity.Core 10.0.6"],
        ["Logging", "Serilog 4.3.1 + Serilog.Sinks.File 7.0.0"],
        ["Visual Basic helper", "Microsoft.VisualBasic 10.3.0"],
    ], [2.2, 3.95], font_size=10.5)

    add_heading(doc, "PHỤ LỤC C. DANH SÁCH ƯU TIÊN SAU REVIEW", 1)
    add_table(doc, ["Ưu tiên", "Công việc", "Tiêu chí hoàn thành"], [
        ["P0", "Di chuyển DDL khỏi luồng nghiệp vụ.", "Không còn ALTER/DROP/CREATE trong Create/Approve/Report."],
        ["P0", "Đối soát BorrowDetails thiếu InstanceID.", "Phiếu hoạt động ánh xạ đúng cá thể hoặc được đánh dấu dữ liệu legacy."],
        ["P1", "Bổ sung integration test mượn/trả.", "Có test cạnh tranh và rollback."],
        ["P1", "Sửa toàn bộ warning.", "Build Release 0 warning, 0 error."],
        ["P1", "Tách ReturnApprovalService.", "Mỗi use case/service có trách nhiệm rõ và test được."],
        ["P2", "Chuẩn hóa role/permission.", "Role ID hoặc permission table thay cho Contains chuỗi."],
        ["P2", "Audit log có cấu trúc.", "Ghi ai, lúc nào, hành động, đối tượng, giá trị trước/sau."],
        ["P2", "Nâng cấp báo cáo.", "Excel có kiểu dữ liệu, style, biểu đồ và PDF."],
    ], [0.75, 2.75, 2.65], font_size=9.8)

    add_heading(doc, "PHỤ LỤC D. CHECKLIST NGHIỆM THU", 1)
    add_table(doc, ["Hạng mục", "Trạng thái tại review", "Ghi chú"], [
        ["Build Release", "Đạt có điều kiện", "0 lỗi, 56 cảnh báo."],
        ["Kết nối SQL Server", "Đạt", "Database QuanLyThietBi truy cập được."],
        ["Khóa ngoại/mã trùng", "Đạt", "Không phát hiện lỗi trong audit chỉ đọc."],
        ["Số lượng cơ bản", "Đạt", "Không có số lượng âm/vượt tổng."],
        ["Dữ liệu cá thể lịch sử", "Chưa đạt", "30 BorrowDetails hoạt động thiếu InstanceID."],
        ["Test tự động", "Chưa đạt", "Không có test project."],
        ["Email nhắc hạn", "Chưa triển khai", "UI đã vô hiệu hóa rõ ràng."],
        ["Tài liệu cài đặt", "Cần bổ sung", "README hiện rất ngắn."],
    ], [2.0, 1.55, 2.6], font_size=10)


def main():
    create_diagrams()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    cover(doc)
    front_matter(doc)
    chapter_intro(doc)
    chapter_requirements(doc)
    chapter_usecase(doc)
    chapter_architecture(doc)
    chapter_database(doc)
    chapter_implementation(doc)
    chapter_review(doc)
    chapter_testing(doc)
    chapter_deployment(doc)
    chapter_future(doc)
    conclusion_appendices(doc)

    core = doc.core_properties
    core.title = "Báo cáo kỹ thuật - Hệ thống quản lý mượn trả thiết bị"
    core.subject = "Dự án Kỹ thuật lập trình"
    core.author = "Bùi Nam Đông"
    core.keywords = "WinForms, .NET 8, SQL Server, quản lý thiết bị, mượn trả"
    core.comments = "Báo cáo được xây dựng từ review mã nguồn và dữ liệu ngày 05/06/2026."
    core.created = datetime(2026, 6, 5)
    core.modified = datetime.now()

    doc.settings.update_fields_on_open = True
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
